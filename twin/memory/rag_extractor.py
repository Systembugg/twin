"""RAG Document Parsing & Semantic Chunking Module.

Supports extracting clean plain text from .pdf, .docx, .txt, .md, .csv, .json, .py
and splitting into semantic chunks for vector database indexing.
"""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)


def extract_text_from_file(file_path: Path) -> str:
    """Extract raw text from various file formats."""
    suffix = file_path.suffix.lower()
    if not file_path.exists():
        return ""

    try:
        if suffix in (".txt", ".md", ".py", ".csv", ".json", ".js", ".ts", ".html", ".css"):
            return file_path.read_text(encoding="utf-8", errors="ignore")

        if suffix == ".docx":
            import docx
            doc = docx.Document(file_path)
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)
            return "\n\n".join(full_text)

        if suffix == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(file_path)
            pages_text = []
            for idx, page in enumerate(reader.pages):
                txt = page.extract_text()
                if txt:
                    pages_text.append(f"--- Page {idx + 1} ---\n{txt}")
            return "\n\n".join(pages_text)

    except Exception as exc:
        log.warning("Could not extract text from file %s: %s", file_path.name, exc)

    return ""


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """Split text into semantic paragraph chunks with overlap."""
    text = text.strip()
    if not text:
        return []

    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    paragraphs = text.split("\n\n")
    current_chunk = ""

    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        if len(current_chunk) + len(p) + 2 <= chunk_size:
            current_chunk = f"{current_chunk}\n\n{p}".strip()
        else:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = p

    if current_chunk:
        chunks.append(current_chunk)

    # Fallback to character chunking if paragraphs are massive
    final_chunks: list[str] = []
    for c in chunks:
        if len(c) <= chunk_size:
            final_chunks.append(c)
        else:
            start = 0
            while start < len(c):
                end = start + chunk_size
                final_chunks.append(c[start:end])
                start += chunk_size - overlap

    return final_chunks
